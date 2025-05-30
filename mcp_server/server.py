"""
LeMCS MCP Server - Model Context Protocol server for legal document processing.

This server exposes LeMCS functionality through standardized MCP tools,
allowing other AI systems to leverage our legal document processing capabilities.
"""
import asyncio
import logging
import os
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime

from mcp.server.fastmcp import FastMCP, Context
from sqlalchemy.ext.asyncio import AsyncSession

from db.database import AsyncSessionLocal
from db.models import Document, Citation, AgentWorkflow, DocumentStatus
from nlp.citation_service import citation_service
from agents.citation_extractor import citation_extractor_agent
from config.settings_simple import settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize MCP server
mcp = FastMCP(
    name="LeMCS Legal Processing Server",
    version="0.1.0",
    description="AI-powered legal document processing and citation analysis"
)


# Document Management Tools

@mcp.tool()
async def upload_document(
    file_path: str,
    document_type: str = "legal_memorandum",
    metadata: Optional[Dict[str, Any]] = None
) -> str:
    """
    Upload a legal document and extract text content.
    
    Args:
        file_path: Path to the document file (DOCX or PDF)
        document_type: Type of document (default: legal_memorandum)
        metadata: Optional metadata dictionary
        
    Returns:
        Document ID of the uploaded document
    """
    try:
        # Validate file exists
        if not Path(file_path).exists():
            raise ValueError(f"File not found: {file_path}")
        
        # Process the document upload directly
        async with AsyncSessionLocal() as db:
            from docx import Document as DocxDocument
            import hashlib
            
            # Calculate file hash
            with open(file_path, 'rb') as f:
                content = f.read()
                content_hash = hashlib.sha256(content).hexdigest()
            
            # Extract text content based on file type
            extracted_text = ""
            filename = Path(file_path).name
            
            if filename.lower().endswith('.docx'):
                try:
                    import io
                    doc = DocxDocument(io.BytesIO(content))
                    extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                except Exception as e:
                    extracted_text = f"Error extracting text: {str(e)}"
            elif filename.lower().endswith('.txt'):
                try:
                    extracted_text = content.decode('utf-8')
                except Exception as e:
                    extracted_text = f"Error reading text file: {str(e)}"
            else:
                extracted_text = "Unsupported file type for text extraction"
            
            # Create document record
            document = Document(
                filename=filename,
                file_path=file_path,
                file_size=len(content),
                content_hash=content_hash,
                extracted_text=extracted_text,
                doc_metadata=metadata or {},
                uploaded_at=datetime.now()
            )
            
            db.add(document)
            await db.commit()
            await db.refresh(document)
            
            return str(document.id)
            
    except Exception as e:
        logger.error(f"Document upload failed: {e}")
        raise RuntimeError(f"Failed to upload document: {str(e)}")


@mcp.tool()
async def get_document(document_id: str) -> Dict[str, Any]:
    """
    Retrieve document metadata and content.
    
    Args:
        document_id: UUID of the document
        
    Returns:
        Dictionary containing document information
    """
    try:
        async with AsyncSessionLocal() as db:
            document = await db.get(Document, document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            return {
                "id": str(document.id),
                "filename": document.filename,
                "file_type": document.file_type,
                "status": document.status.value,
                "extracted_text_preview": document.extracted_text[:500] + "..." if document.extracted_text and len(document.extracted_text) > 500 else document.extracted_text,
                "uploaded_at": document.uploaded_at.isoformat() if document.uploaded_at else None,
                "processed_at": document.processed_at.isoformat() if document.processed_at else None,
                "metadata": document.doc_metadata
            }
            
    except Exception as e:
        logger.error(f"Failed to get document {document_id}: {e}")
        raise RuntimeError(f"Failed to retrieve document: {str(e)}")


@mcp.tool()
async def list_documents(
    status: Optional[str] = None,
    limit: int = 10
) -> List[Dict[str, Any]]:
    """
    List available documents with optional filtering.
    
    Args:
        status: Optional status filter (uploaded, processing, processed, failed)
        limit: Maximum number of documents to return (default: 10)
        
    Returns:
        List of document summaries
    """
    try:
        from sqlalchemy import select, desc
        
        async with AsyncSessionLocal() as db:
            query = select(Document).order_by(desc(Document.uploaded_at)).limit(limit)
            
            # Add status filter if provided
            if status:
                try:
                    status_enum = DocumentStatus(status.lower())
                    query = query.where(Document.status == status_enum)
                except ValueError:
                    raise ValueError(f"Invalid status: {status}. Valid statuses: uploaded, processing, processed, failed")
            
            result = await db.execute(query)
            documents = result.scalars().all()
            
            return [
                {
                    "id": str(doc.id),
                    "filename": doc.filename,
                    "document_type": doc.document_type.value,
                    "status": doc.status.value,
                    "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
                    "has_text": bool(doc.extracted_text)
                }
                for doc in documents
            ]
            
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        raise RuntimeError(f"Failed to list documents: {str(e)}")


# Citation Extraction Tools

@mcp.tool()
async def extract_citations(
    document_id: str,
    resolve_references: bool = True,
    analyze_authority: bool = True
) -> Dict[str, Any]:
    """
    Extract and analyze legal citations from a document.
    
    Args:
        document_id: UUID of the document to process
        resolve_references: Whether to resolve supra/id citations (default: True)
        analyze_authority: Whether to analyze citation authority (default: True)
        
    Returns:
        Dictionary containing extracted citations and analysis results
    """
    try:
        async with AsyncSessionLocal() as db:
            # Get the document
            document = await db.get(Document, document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            if not document.extracted_text:
                raise ValueError("Document has no extracted text content")
            
            # Process options
            options = {
                "resolve_references": resolve_references,
                "create_embeddings": False,  # Not implemented yet
                "analyze_authority": analyze_authority
            }
            
            # Run citation extraction workflow
            result = await citation_extractor_agent.process_document(
                document=document,
                db_session=db,
                options=options
            )
            
            # Format response
            citations = []
            authority_analysis = {a.citation.id: a for a in result.get("authority_analysis", [])}
            
            for citation in result.get("citations", []):
                analysis = authority_analysis.get(citation.id)
                citation_data = {
                    "id": str(citation.id),
                    "text": citation.citation_text,
                    "type": citation.citation_type,
                    "reporter": citation.reporter,
                    "volume": citation.volume,
                    "page": citation.page,
                    "position": [citation.position_start, citation.position_end],
                    "confidence_score": citation.confidence_score
                }
                
                if analysis:
                    citation_data.update({
                        "court_level": analysis.court_level,
                        "jurisdiction": analysis.jurisdiction,
                        "precedential_strength": analysis.precedential_strength,
                        "authority_score": analysis.authority_score
                    })
                
                citations.append(citation_data)
            
            return {
                "document_id": document_id,
                "citations": citations,
                "extraction_stats": result.get("extraction_result", {}).extraction_stats or {},
                "workflow_id": str(result.get("workflow_id", "")),
                "processing_time_ms": result.get("metrics", {}).get("extraction_time_ms", 0),
                "errors": result.get("errors", [])
            }
            
    except Exception as e:
        logger.error(f"Citation extraction failed for {document_id}: {e}")
        raise RuntimeError(f"Citation extraction failed: {str(e)}")


@mcp.tool()
async def get_document_citations(
    document_id: str,
    include_authority: bool = True
) -> List[Dict[str, Any]]:
    """
    Get all citations for a specific document.
    
    Args:
        document_id: UUID of the document
        include_authority: Whether to include authority analysis (default: True)
        
    Returns:
        List of citations with optional authority analysis
    """
    try:
        from sqlalchemy import select
        
        async with AsyncSessionLocal() as db:
            # Validate document exists
            document = await db.get(Document, document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            # Get citations from database
            query = select(Citation).where(Citation.document_id == document_id)
            result = await db.execute(query)
            citations = result.scalars().all()
            
            # Format response
            citation_responses = []
            for citation in citations:
                citation_data = {
                    "id": str(citation.id),
                    "text": citation.citation_text,
                    "type": citation.citation_type,
                    "reporter": citation.reporter,
                    "volume": citation.volume,
                    "page": citation.page,
                    "position": [citation.position_start, citation.position_end],
                    "confidence_score": citation.confidence_score
                }
                
                # Add authority analysis if requested
                if include_authority:
                    try:
                        analysis = await citation_service.analyze_citation_authority(citation)
                        citation_data.update({
                            "court_level": analysis.court_level,
                            "jurisdiction": analysis.jurisdiction,
                            "precedential_strength": analysis.precedential_strength,
                            "authority_score": analysis.authority_score
                        })
                    except Exception as e:
                        logger.warning(f"Authority analysis failed for citation {citation.id}: {e}")
                
                citation_responses.append(citation_data)
            
            return citation_responses
            
    except Exception as e:
        logger.error(f"Failed to get citations for document {document_id}: {e}")
        raise RuntimeError(f"Failed to retrieve citations: {str(e)}")


@mcp.tool()
async def search_citations(
    query: str,
    document_ids: Optional[List[str]] = None,
    limit: int = 10
) -> List[Dict[str, Any]]:
    """
    Search citations by text content.
    
    Args:
        query: Search query for citations
        document_ids: Optional list of document IDs to limit search
        limit: Maximum number of results (default: 10)
        
    Returns:
        List of matching citations
    """
    try:
        from sqlalchemy import select, or_
        
        async with AsyncSessionLocal() as db:
            # Build search query
            query_obj = select(Citation)
            
            # Add text search filter
            search_filters = [
                Citation.citation_text.ilike(f"%{query}%"),
                Citation.reporter.ilike(f"%{query}%")
            ]
            query_obj = query_obj.where(or_(*search_filters))
            
            # Add document filter if specified
            if document_ids:
                query_obj = query_obj.where(Citation.document_id.in_(document_ids))
            
            # Add limit
            query_obj = query_obj.limit(limit)
            
            # Execute query
            result = await db.execute(query_obj)
            citations = result.scalars().all()
            
            # Format response
            return [
                {
                    "id": str(citation.id),
                    "text": citation.citation_text,
                    "type": citation.citation_type,
                    "reporter": citation.reporter,
                    "volume": citation.volume,
                    "page": citation.page,
                    "document_id": str(citation.document_id),
                    "confidence_score": citation.confidence_score
                }
                for citation in citations
            ]
            
    except Exception as e:
        logger.error(f"Citation search failed: {e}")
        raise RuntimeError(f"Citation search failed: {str(e)}")


# Analysis Tools

@mcp.tool()
async def analyze_citation_authority(citation_id: str) -> Dict[str, Any]:
    """
    Analyze the precedential authority of a specific citation.
    
    Args:
        citation_id: UUID of the citation
        
    Returns:
        Authority analysis results
    """
    try:
        async with AsyncSessionLocal() as db:
            citation = await db.get(Citation, citation_id)
            if not citation:
                raise ValueError(f"Citation {citation_id} not found")
            
            analysis = await citation_service.analyze_citation_authority(citation)
            
            return {
                "citation_id": str(citation.id),
                "citation_text": citation.citation_text,
                "court_level": analysis.court_level,
                "jurisdiction": analysis.jurisdiction,
                "precedential_strength": analysis.precedential_strength,
                "authority_score": analysis.authority_score,
                "confidence_score": analysis.confidence_score
            }
            
    except Exception as e:
        logger.error(f"Authority analysis failed for citation {citation_id}: {e}")
        raise RuntimeError(f"Authority analysis failed: {str(e)}")


@mcp.tool()
async def get_citation_statistics(
    document_id: Optional[str] = None
) -> Dict[str, Any]:
    """
    Get citation extraction performance statistics.
    
    Args:
        document_id: Optional document ID for document-specific stats
        
    Returns:
        Statistics about citation extraction performance
    """
    try:
        from sqlalchemy import select, func
        
        async with AsyncSessionLocal() as db:
            # Get service statistics
            service_stats = citation_service.get_extraction_statistics()
            
            # Get database statistics
            if document_id:
                # Document-specific stats
                query = select(func.count(Citation.id)).where(Citation.document_id == document_id)
                result = await db.execute(query)
                citation_count = result.scalar() or 0
                
                stats = {
                    "document_id": document_id,
                    "total_citations": citation_count,
                    **service_stats
                }
            else:
                # Overall stats
                query = select(func.count(Citation.id))
                result = await db.execute(query)
                total_citations = result.scalar() or 0
                
                query = select(func.count(Document.id.distinct())).select_from(
                    Citation.__table__.join(Document.__table__)
                )
                result = await db.execute(query)
                documents_with_citations = result.scalar() or 0
                
                stats = {
                    "total_citations": total_citations,
                    "documents_with_citations": documents_with_citations,
                    **service_stats
                }
            
            return stats
            
    except Exception as e:
        logger.error(f"Failed to get citation statistics: {e}")
        raise RuntimeError(f"Failed to retrieve statistics: {str(e)}")


# Workflow Management Tools

@mcp.tool()
async def get_workflows(
    agent_type: Optional[str] = None,
    limit: int = 10
) -> List[Dict[str, Any]]:
    """
    Get recent agent workflow executions.
    
    Args:
        agent_type: Optional agent type filter (citation_extractor, etc.)
        limit: Maximum number of workflows (default: 10)
        
    Returns:
        List of workflow execution information
    """
    try:
        from sqlalchemy import select, desc
        from db.models import AgentType
        
        async with AsyncSessionLocal() as db:
            # Get recent workflows
            query = (
                select(AgentWorkflow)
                .order_by(desc(AgentWorkflow.created_at))
                .limit(limit)
            )
            
            # Add agent type filter if provided
            if agent_type:
                try:
                    agent_type_enum = AgentType(agent_type.upper())
                    query = query.where(AgentWorkflow.agent_type == agent_type_enum)
                except ValueError:
                    raise ValueError(f"Invalid agent type: {agent_type}")
            
            result = await db.execute(query)
            workflows = result.scalars().all()
            
            # Format response
            return [
                {
                    "id": str(workflow.id),
                    "agent_type": workflow.agent_type.value,
                    "status": workflow.status.value,
                    "started_at": workflow.started_at.isoformat() if workflow.started_at else None,
                    "completed_at": workflow.completed_at.isoformat() if workflow.completed_at else None,
                    "execution_time_ms": workflow.execution_time_ms,
                    "quality_score": workflow.quality_score,
                    "retry_count": workflow.retry_count,
                    "error_message": workflow.error_message
                }
                for workflow in workflows
            ]
            
    except Exception as e:
        logger.error(f"Failed to get workflows: {e}")
        raise RuntimeError(f"Failed to retrieve workflows: {str(e)}")


@mcp.tool()
async def get_workflow_status(workflow_id: str) -> Dict[str, Any]:
    """
    Get detailed status of a specific workflow.
    
    Args:
        workflow_id: UUID of the workflow
        
    Returns:
        Detailed workflow information
    """
    try:
        async with AsyncSessionLocal() as db:
            workflow = await db.get(AgentWorkflow, workflow_id)
            if not workflow:
                raise ValueError(f"Workflow {workflow_id} not found")
            
            return {
                "id": str(workflow.id),
                "agent_type": workflow.agent_type.value,
                "status": workflow.status.value,
                "started_at": workflow.started_at.isoformat() if workflow.started_at else None,
                "completed_at": workflow.completed_at.isoformat() if workflow.completed_at else None,
                "execution_time_ms": workflow.execution_time_ms,
                "quality_score": workflow.quality_score,
                "retry_count": workflow.retry_count,
                "input_data": workflow.input_data,
                "output_data": workflow.output_data,
                "error_message": workflow.error_message
            }
            
    except Exception as e:
        logger.error(f"Failed to get workflow {workflow_id}: {e}")
        raise RuntimeError(f"Failed to retrieve workflow: {str(e)}")


# Resources

@mcp.resource("document://{document_id}")
async def get_document_content(document_id: str) -> str:
    """Get the full text content of a document."""
    try:
        async with AsyncSessionLocal() as db:
            document = await db.get(Document, document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            if not document.extracted_text:
                return f"Document {document_id} has no extracted text content"
            
            return document.extracted_text
            
    except Exception as e:
        logger.error(f"Failed to get document content for {document_id}: {e}")
        return f"Error retrieving document content: {str(e)}"


@mcp.resource("citations://{document_id}")
async def get_document_citations_resource(document_id: str) -> str:
    """Get all citations for a document as formatted text."""
    try:
        citations = await get_document_citations(document_id, include_authority=True)
        
        if not citations:
            return f"No citations found for document {document_id}"
        
        # Format citations as readable text
        formatted_citations = []
        for citation in citations:
            authority_info = ""
            if citation.get("authority_score"):
                authority_info = f" (Authority: {citation['precedential_strength']}, Score: {citation['authority_score']:.2f})"
            
            formatted_citations.append(
                f"- {citation['text']}{authority_info}"
            )
        
        return f"Citations for document {document_id}:\n" + "\n".join(formatted_citations)
        
    except Exception as e:
        logger.error(f"Failed to get citations resource for {document_id}: {e}")
        return f"Error retrieving citations: {str(e)}"


def main():
    """Main entry point for the MCP server."""
    try:
        # Validate required environment
        if not settings.DATABASE_URL:
            raise RuntimeError("DATABASE_URL not configured")
        
        logger.info("Starting LeMCS MCP Server...")
        logger.info(f"Database URL: {settings.DATABASE_URL}")
        
        # Run the server with stdio transport
        mcp.run(transport="stdio")
        
    except KeyboardInterrupt:
        logger.info("Server shutdown requested")
    except Exception as e:
        logger.error(f"Server startup failed: {e}")
        raise


if __name__ == "__main__":
    main()